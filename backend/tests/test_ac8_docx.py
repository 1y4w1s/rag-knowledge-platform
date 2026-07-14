"""AC-8：DOCX 上传 → 入库 → 对话检索 → 引用溯源（PRD §11）。"""

from __future__ import annotations

import json
import re
import uuid
from pathlib import Path

import pytest
from httpx import AsyncClient
from sqlalchemy import select

from app.core.config import settings
from app.core.database import SessionLocal
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.models.enums import DocumentStatus
from app.services.ingestion.pipeline import process_document_ingestion
from tests.conftest import create_test_kb as _create_kb

FIXTURES = Path(__file__).parent / "fixtures"
GOLDEN_DOCX = FIXTURES / "golden_handbook.docx"


def _make_golden_docx(path: Path) -> None:
    """与 golden_handbook.md 对齐的最小 DOCX（Heading 样式 + 年假条款）。"""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("员工手册", level=1)
    doc.add_heading("第一章 考勤制度", level=2)
    doc.add_heading("1.1 年假", level=3)
    doc.add_paragraph(
        "员工年满一年后可享受年假10天。年假须提前两周申请，由直属主管审批后方可休假。"
    )
    doc.add_heading("1.2 迟到", level=3)
    doc.add_paragraph("迟到 30 分钟以内按事假半天处理；超过 30 分钟按旷工半天处理。")
    doc.save(str(path))


def _parse_sse_events(raw: str) -> list[tuple[str, dict]]:
    events: list[tuple[str, dict]] = []
    for block in re.split(r"\n\n+", raw.strip()):
        if not block.strip():
            continue
        event_name = "message"
        data_str = ""
        for line in block.splitlines():
            if line.startswith("event: "):
                event_name = line.removeprefix("event: ").strip()
            elif line.startswith("data: "):
                data_str = line.removeprefix("data: ").strip()
        if data_str:
            events.append((event_name, json.loads(data_str)))
    return events


async def _ingest_fixture(
    *,
    kb_id: uuid.UUID,
    user_id: uuid.UUID,
    source: Path,
    file_type: str,
    upload_dir: Path,
) -> Document:
    doc_id = uuid.uuid4()
    storage_dir = upload_dir / str(kb_id) / str(doc_id)
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_path = storage_dir / f"{uuid.uuid4()}.{file_type}"
    storage_path.write_bytes(source.read_bytes())

    async with SessionLocal() as db:
        doc = Document(
            id=doc_id,
            kb_id=kb_id,
            filename=source.name,
            file_type=file_type,
            file_size=storage_path.stat().st_size,
            storage_path=str(storage_path),
            status=DocumentStatus.queued,
            uploaded_by=user_id,
        )
        db.add(doc)
        await db.commit()

    await process_document_ingestion(doc_id)

    async with SessionLocal() as db:
        row = await db.get(Document, doc_id)
        assert row is not None
        assert row.status == DocumentStatus.completed
        return row


async def _load_chunks(document_id: uuid.UUID) -> list[DocumentChunk]:
    async with SessionLocal() as db:
        result = await db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index)
        )
        return list(result.all())


@pytest.fixture(scope="session")
def golden_docx_fixture() -> Path:
    """确保 committed fixture 存在；缺失时按 golden md 内容生成。"""
    if not GOLDEN_DOCX.exists():
        _make_golden_docx(GOLDEN_DOCX)
    return GOLDEN_DOCX


@pytest.fixture
def upload_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))
    return tmp_path


@pytest.mark.asyncio
async def test_ac8_docx_ingestion_chunks_have_section_metadata(
    client: AsyncClient,
    register_and_login,
    upload_dir: Path,
    golden_docx_fixture: Path,
) -> None:
    """AC-8 前半：DOCX 入库后切片含章节元数据。"""
    headers, user = await register_and_login(prefix="ac8-docx-ingest")
    kb = await _create_kb(client, headers, user, name="AC8 DOCX 库")
    kb_id = uuid.UUID(kb["id"])
    user_id = uuid.UUID(user["id"])

    doc = await _ingest_fixture(
        kb_id=kb_id,
        user_id=user_id,
        source=golden_docx_fixture,
        file_type="docx",
        upload_dir=upload_dir,
    )

    assert doc.chunk_count is not None and doc.chunk_count > 0
    chunks = await _load_chunks(doc.id)
    annual = next(c for c in chunks if "年假10天" in c.content)
    assert annual.section_title == "1.1 年假"
    assert annual.heading_path is not None
    assert "考勤制度" in annual.heading_path


@pytest.mark.asyncio
async def test_ac8_docx_chat_returns_citation(
    client: AsyncClient,
    register_and_login,
    upload_dir: Path,
    golden_docx_fixture: Path,
) -> None:
    """AC-8：上传 DOCX → 提问 → SSE 引用含文档名、章节与年假片段。"""
    headers, user = await register_and_login(prefix="ac8-docx-chat")
    kb = await _create_kb(client, headers, user, name="AC8 对话库")
    kb_id = uuid.UUID(kb["id"])
    user_id = uuid.UUID(user["id"])

    await _ingest_fixture(
        kb_id=kb_id,
        user_id=user_id,
        source=golden_docx_fixture,
        file_type="docx",
        upload_dir=upload_dir,
    )

    async with client.stream(
        "POST",
        f"/api/v1/knowledge-bases/{kb_id}/chat",
        headers=headers,
        json={"message": "员工年假有几天？"},
    ) as resp:
        body = await resp.aread()
        assert resp.status_code == 200
        events = _parse_sse_events(body.decode("utf-8"))

    citations = [data for name, data in events if name == "citation"]
    assert citations, "应返回至少一条 citation SSE 事件"
    assert all(c.get("doc_name") == golden_docx_fixture.name for c in citations)
    assert any(
        "年假" in c.get("excerpt", "") or "10" in c.get("excerpt", "") for c in citations
    )
    assert any(c.get("section_title") == "1.1 年假" for c in citations)

    done = next(data for name, data in events if name == "done")
    assert len(done.get("citations", [])) >= 1
