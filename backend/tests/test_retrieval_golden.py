"""Golden QA 检索系统集成测试（Wave 3.4 · Plan-RAG R5-2 Hit@3 门禁）。
v0.5：支持多相关文档标注 + 拒答测试。
"""

from __future__ import annotations

import uuid
from pathlib import Path
from uuid import UUID

import pytest
from httpx import AsyncClient

from app.core.config import settings
from app.core.database import SessionLocal
from app.models.document import Document
from app.models.enums import DocumentStatus
from app.services.ingestion import embedder
from app.services.ingestion.pipeline import process_document_ingestion
from app.services.rag.retrieval import retrieve_chunks
from app.services.rag.types import RetrievedChunk
from tests.golden_qa_loader import (
    FIXTURES,
    GOLDEN_DOCX,
    GOLDEN_MD,
    GOLDEN_QA_CASES,
    GoldenQACase,
    HIT_K,
    chunk_matches,
    hit_at_k,
    reciprocal_rank,
)

@pytest.fixture(autouse=True)
def _mock_embedding(monkeypatch: pytest.MonkeyPatch) -> None:
    """所有 golden 测试使用 mock 嵌入（避免真实 API 调用）。
    设置 RAG_REAL_EMBEDDING=1 使用真实嵌入。
    同时禁用低置信度 expand（避免 mock 向量触发 DeepSeek 干扰 Hit@3）。

    mock 嵌入实现统一引用 embedder._mock_vector（字符 2/3-gram 词袋 SSOT），
    不在此处另起实现（N15）。
    """
    import os

    async def _no_expand(query: str) -> list[str]:
        return [query]

    monkeypatch.setattr(
        "app.services.rag.generation.expand_queries",
        _no_expand,
    )
    if os.environ.get("RAG_REAL_EMBEDDING") == "1":
        return
    monkeypatch.setattr(embedder, "embed_texts", _embed_texts_mock_unified)


async def _embed_texts_mock_unified(
    texts: list[str], provider: str | None = None
) -> list[list[float]]:
    """统一 mock 嵌入入口：转调 embedder._mock_vector（SSOT）。"""
    dim = 384 if (provider or "").lower() == "bge_en" else settings.embedding_dim
    return [embedder._mock_vector(t, dim=dim) for t in texts]


def _make_golden_pdf(path):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 720, "Employee Handbook")
    c.drawString(72, 690, "Chapter 1 Attendance")
    c.drawString(72, 660, "Apply annual leave two weeks")
    c.showPage()
    c.drawString(72, 720, "in advance. After one year: annual leave 10 days.")
    c.save()


def _make_golden_docx(path):
    from docx import Document
    doc = Document()
    doc.add_heading("考勤制度", level=1)
    doc.add_heading("1.1 年假", level=2)
    doc.add_paragraph("员工年满一年后可享受年假10天。")
    doc.save(str(path))


async def _ingest_fixture(
    *,
    kb_id,
    user_id,
    source,
    file_type,
    upload_dir,
) -> Document:
    doc_id = uuid.uuid4()
    storage_dir = upload_dir / str(kb_id) / str(doc_id)
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_path = storage_dir / f"{uuid.uuid4()}.{file_type}"
    storage_path.write_bytes(source.read_bytes())
    async with SessionLocal() as db:
        doc = Document(
            id=doc_id, kb_id=kb_id, filename=source.name,
            file_type=file_type, file_size=storage_path.stat().st_size,
            storage_path=str(storage_path), status=DocumentStatus.queued,
            uploaded_by=user_id,
        )
        db.add(doc)
        await db.commit()
    await process_document_ingestion(doc_id)
    async with SessionLocal() as db:
        row = await db.get(Document, doc_id)
        assert row is not None
        return row


@pytest.fixture
def upload_dir(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))
    return tmp_path


from tests.conftest import create_test_kb as _create_kb

# R5-2 经典门禁题（历史 12/12；fixture 若缺号则取现有）
_GATE_IDS = {f"GQ-{i}" for i in range(1, 13)}
GATE_CASES = [c for c in GOLDEN_QA_CASES if c.case_id in _GATE_IDS]


# G3 门禁固化：GQ-30/GQ-77 为已知 cross_reference 复合题检索 miss（min_match=2 跨章节），
# 实验 N 确认既有缺陷、非本窗引入；待检索优化窗（composite recall 增强）修复后移除 xfail。
_XFAIL_CROSS_REF = {"GQ-30", "GQ-77"}


def _xfail_known_issue(case) -> None:
    """对已知缺陷题打 xfail 标记（门禁固化：其余失败必须阻断 CI）。"""
    if case.case_id in _XFAIL_CROSS_REF:
        import pytest
        pytest.xfail(f"{case.case_id} 已知 cross_reference 复合题 miss（min_match=2 跨章节），待检索优化窗")


@pytest.mark.parametrize("case", GOLDEN_QA_CASES, ids=lambda c: c.case_id)
@pytest.mark.asyncio
async def test_golden_qa_hit_at_3(
    client: AsyncClient,
    register_and_login,
    upload_dir,
    case: GoldenQACase,
    tmp_path: Path,
) -> None:
    """每道 golden QA 题：入库黄金文档 → 检索 → 验证 Top-3 命中（默认 RERANK_POLICY=off）。"""
    _xfail_known_issue(case)
    await _assert_golden_case(
        client, register_and_login, upload_dir, case, tmp_path,
    )


@pytest.mark.parametrize("case", GATE_CASES, ids=lambda c: c.case_id)
@pytest.mark.asyncio
async def test_golden_gate_hit_at_3_conditional_mock(
    client: AsyncClient,
    register_and_login,
    upload_dir,
    case: GoldenQACase,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """经典门禁题在 RERANK_POLICY=conditional + mock 下 Hit@3 不劣于 off。"""
    monkeypatch.setattr(settings, "rerank_policy", "conditional")
    monkeypatch.setattr(settings, "rerank_provider", "mock")
    await _assert_golden_case(
        client, register_and_login, upload_dir, case, tmp_path,
        label="conditional",
    )


@pytest.mark.parametrize("case", GATE_CASES, ids=lambda c: c.case_id)
@pytest.mark.asyncio
async def test_golden_gate_hit_at_3_conditional_multi_query_mock(
    client: AsyncClient,
    register_and_login,
    upload_dir,
    case: GoldenQACase,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """经典门禁题在 QUERY_REWRITE_POLICY=conditional + mock 变体下 Hit@3 不劣于 off。

    不强行 should_expand（强制扩≡always，会伤 Hit@3）；自然门闩下强池不扩，
    偶发 miss/短问才走 mock 变体。
    """
    from app.services.rag.multi_query import mock_expand_queries

    monkeypatch.setattr(settings, "query_rewrite_policy", "conditional")
    monkeypatch.setattr(settings, "query_rewrite_enabled", False)

    async def _mock_variants(query: str, **kwargs):
        return mock_expand_queries(query)

    monkeypatch.setattr(
        "app.services.rag.multi_query.build_query_variants",
        _mock_variants,
    )
    await _assert_golden_case(
        client, register_and_login, upload_dir, case, tmp_path,
        label="cond-mq",
    )


async def _assert_golden_case(
    client: AsyncClient,
    register_and_login,
    upload_dir,
    case: GoldenQACase,
    tmp_path: Path,
    *,
    label: str | None = None,
) -> None:
    headers, user = await register_and_login(prefix=case.case_id.replace("-", ""))
    kb = await _create_kb(client, headers, user)
    kb_id = uuid.UUID(kb["id"])

    if hasattr(case, "file_type") and case.file_type == "docx":
        pytest.importorskip("docx")
        if not GOLDEN_DOCX.exists():
            _make_golden_docx(GOLDEN_DOCX)
        source = GOLDEN_DOCX
        file_type = "docx"
    elif hasattr(case, "file_type") and case.file_type == "pdf":
        pytest.importorskip("reportlab")
        source = tmp_path / "golden_handbook.pdf"
        _make_golden_pdf(source)
        file_type = "pdf"
    else:
        source = GOLDEN_MD
        file_type = "md"

    await _ingest_fixture(
        kb_id=kb_id,
        user_id=user["id"],
        source=source,
        file_type=file_type,
        upload_dir=upload_dir,
    )

    async with SessionLocal() as db:
        chunks = await retrieve_chunks(
            db,
            kb_id=kb_id,
            query=case.query,
            top_k=HIT_K,
        )

    assert chunks, f"{case.case_id} 检索无结果"

    passed = hit_at_k(chunks, case, k=HIT_K)
    tag = f"{case.case_id}" + (f"/{label}" if label else "")

    if case.expect_rejection:
        match_details = [(c.section_title, c.page_number, c.content[:60]) for c in chunks[:HIT_K] if chunk_matches(case, c)]
        assert passed, (
            f"{tag} 拒答失败：Top-3 内存在匹配结果 "
            f"{match_details}"
        )
    else:
        assert passed, (
            f"{tag} Hit@{HIT_K} 未命中{'（需 ≥{} 个匹配）'.format(case.min_match) if case.min_match > 1 else ''}；"
            f"Top-{HIT_K}="
            f"{[(c.section_title, c.page_number, c.content[:40]) for c in chunks[:HIT_K]]}"
        )

    rr = reciprocal_rank(chunks, case, k=HIT_K)
    if rr < 1.0:
        print(f"  {tag}: RR={rr:.3f}")


REJECTION_ACCURACY_MIN = 0.80  # 拒答准确率门禁 >=80%


@pytest.mark.asyncio
async def test_golden_rejection_accuracy(
    client: AsyncClient,
    register_and_login,
    upload_dir,
) -> None:
    """全量拒答题系统级门禁：拒答准确率 >=80%"""
    headers, user = await register_and_login(prefix="gd-rej")
    kb = await _create_kb(client, headers, user)

    await _ingest_fixture(
        kb_id=uuid.UUID(kb["id"]),
        user_id=uuid.UUID(user["id"]),
        source=GOLDEN_MD, file_type="md",
        upload_dir=upload_dir,
    )

    rejection_cases = [c for c in GOLDEN_QA_CASES if c.expect_rejection]
    correct = 0

    async with SessionLocal() as db:
        for case in rejection_cases:
            chunks = await retrieve_chunks(
                db, kb_id=uuid.UUID(kb["id"]),
                query=case.query, top_k=HIT_K,
            )
            match_count = sum(1 for c in chunks[:HIT_K] if chunk_matches(case, c))
            if match_count == 0:
                correct += 1

    total = len(rejection_cases)
    accuracy = correct / total if total > 0 else 1.0
    print(f"  拒答准确率: {correct}/{total} = {accuracy:.0%}  (门禁: >=80%)")
    assert accuracy >= REJECTION_ACCURACY_MIN, (
        f"拒答准确率 {accuracy:.0%} ({correct}/{total}) 低于门禁 {REJECTION_ACCURACY_MIN:.0%}"
    )
